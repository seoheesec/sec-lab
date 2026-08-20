from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "SecureLens_최종보고서_제안서기반.docx"
TARGET = ROOT / "outputs" / "SecureLens_최종보고서_최종.docx"


SCHEDULE_TEXT = (
    "개발 일정은 1주차에 주제 선정과 제안서 작성을 진행하고, 2~3주차에는 로그인, "
    "대시보드, 분석 화면 등 주요 UI를 구현하는 방식으로 진행하였다. 4주차에는 정적 분석과 "
    "취약점 카드 등 핵심 기능을 구현하였고, 5주차에는 기존 프로젝트의 실행 오류를 복구하면서 "
    "Codex와 함께 Protected Route, GitHub Connect, AI 분석, 오탐 검토, 사용자별 저장 구조를 "
    "재구현하였다. 6주차에는 오탐 수 반영, Python 언어 감지, 마이페이지 통계 오류를 수정하며 "
    "테스트를 진행하였고, 7주차에는 최종 보고서와 발표 자료, 코드 설명 자료를 정리하였다."
)


def paragraph_text(element):
    return "".join(t.text or "" for t in element.xpath(".//w:t"))


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    # Copy paragraph style/spacing from a normal body paragraph near the source.
    new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = new_para.paragraph_format
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.15
    run = new_para.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)
    return new_para


def remove_from_marker(doc, marker):
    body = doc._body._element
    children = list(body)
    start = None
    for idx, child in enumerate(children):
        if child.tag == qn("w:p") and paragraph_text(child).strip().startswith(marker):
            start = idx
            break
    if start is None:
        return 0
    removed = 0
    for child in children[start:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
        removed += 1
    return removed


def ensure_korean_font(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.font.name = run.font.name or "Malgun Gothic"
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.rFonts
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:eastAsia"), "Malgun Gothic")


def replace_cell_text(doc):
    replacements = {
        "Monaco Editor, 파일 업로드, 폴더 업로드, GitHub 저장소 연결 기능을 제공한다. 실제 제출 전 실행 화면 캡처를 삽입한다.": (
            "Monaco Editor, 파일 업로드, 폴더 업로드, GitHub 저장소 연결 기능을 제공하여 사용자가 분석할 코드를 직접 확인한 뒤 제출할 수 있게 하였다."
        ),
        "보안 점수, 위험도 분포, 취약점 목록, 취약점 상세 모달로 이어지는 구조를 제공한다. 실제 제출 전 화면 캡처를 삽입한다.": (
            "보안 점수, 위험도 분포, 취약점 목록, 취약점 상세 모달로 이어지는 구조를 제공하여 분석 결과를 한눈에 확인할 수 있게 하였다."
        ),
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                current = cell.text
                if current in replacements:
                    cell.text = replacements[current]


def main():
    doc = Document(SOURCE)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "1-6 개발일정":
            next_el = paragraph._p.getnext()
            if not (next_el is not None and SCHEDULE_TEXT[:20] in paragraph_text(next_el)):
                insert_paragraph_after(paragraph, SCHEDULE_TEXT)
            break

    remove_from_marker(doc, "부록")
    replace_cell_text(doc)
    ensure_korean_font(doc)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
